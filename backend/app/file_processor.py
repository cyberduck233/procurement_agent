"""
文件处理模块 - 支持 PDF、Word、TXT、图片等文件的解析和向量化
"""
import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import hashlib
from datetime import datetime

logger = logging.getLogger(__name__)


class FileProcessor:
    """文件处理器 - 解析各种格式文件并提取文本"""
    
    def __init__(self, upload_dir: str = "data/uploads"):
        # 使用相对于当前文件的路径
        current_dir = Path(__file__).parent.parent  # backend/app -> backend
        self.upload_dir = current_dir / upload_dir
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"📁 文件上传目录: {self.upload_dir.absolute()}")
        
    def save_file(self, file_content: bytes, filename: str) -> str:
        """保存上传的文件"""
        # 生成安全的文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = f"{timestamp}_{filename}"
        file_path = self.upload_dir / safe_filename
        
        with open(file_path, "wb") as f:
            f.write(file_content)
        
        logger.info(f"✅ 文件已保存: {file_path}")
        return str(file_path)
    
    def extract_text(self, file_path: str) -> str:
        """从文件中提取文本内容"""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        try:
            if extension == ".txt":
                return self._extract_from_txt(file_path)
            elif extension == ".pdf":
                return self._extract_from_pdf(file_path)
            elif extension in [".doc", ".docx"]:
                return self._extract_from_word(file_path)
            elif extension in [".jpg", ".jpeg", ".png", ".bmp"]:
                return self._extract_from_image(file_path)
            else:
                logger.warning(f"⚠️ 不支持的文件类型: {extension}")
                return f"[不支持的文件类型: {extension}]"
        except Exception as e:
            logger.error(f"❌ 文件解析失败 {file_path}: {e}")
            return f"[文件解析失败: {str(e)}]"
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """提取 TXT 文件内容"""
        try:
            # 尝试多种编码
            for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    logger.info(f"✅ 成功读取 TXT (编码: {encoding})")
                    return content
                except UnicodeDecodeError:
                    continue
            
            logger.warning("⚠️ 所有编码尝试失败，使用 latin-1 强制读取")
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            raise Exception(f"TXT 文件读取失败: {e}")
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """提取 PDF 文件内容"""
        try:
            import PyPDF2
            
            text_parts = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)
                
                logger.info(f"📖 开始提取 PDF，共 {num_pages} 页")
                
                for page_num in range(min(num_pages, 50)):  # 最多提取前50页，避免太长
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)
                
                if num_pages > 50:
                    logger.warning(f"⚠️ PDF 太长，仅提取前 50 页（共 {num_pages} 页）")
            
            result = "\n\n".join(text_parts)
            logger.info(f"✅ 成功提取 PDF，共 {min(num_pages, 50)} 页，{len(result)} 字符")
            return result if result else "[PDF 文件为空或无法提取文本]"
            
        except ImportError:
            logger.error("❌ PyPDF2 未安装，请运行: pip install PyPDF2")
            return "[需要安装 PyPDF2: pip install PyPDF2]"
        except Exception as e:
            logger.error(f"❌ PDF 解析异常: {e}")
            raise Exception(f"PDF 解析失败: {e}")
    
    def _extract_from_word(self, file_path: Path) -> str:
        """提取 Word 文件内容"""
        try:
            import docx
            
            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # 提取表格
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        tables_text.append(row_text)
            
            content_parts = []
            if paragraphs:
                content_parts.append("\n".join(paragraphs))
            if tables_text:
                content_parts.append("\n--- 表格内容 ---\n" + "\n".join(tables_text))
            
            result = "\n\n".join(content_parts)
            logger.info(f"✅ 成功提取 Word，{len(paragraphs)} 段落，{len(result)} 字符")
            return result if result else "[Word 文件为空]"
            
        except ImportError:
            logger.error("❌ python-docx 未安装，请运行: pip install python-docx")
            return "[需要安装 python-docx: pip install python-docx]"
        except Exception as e:
            raise Exception(f"Word 解析失败: {e}")
    
    def _extract_from_image(self, file_path: Path) -> str:
        """提取图片中的文字（OCR）"""
        logger.warning("⚠️ 图片 OCR 功能需要 pytesseract，暂不支持")
        return f"[图片文件: {file_path.name}，OCR 功能待实现]"
    
    def calculate_file_hash(self, file_path: str) -> str:
        """计算文件 MD5 哈希值，用于去重"""
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """将长文本分割成小块"""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # 尝试在句号、换行符等处分割
        if end < len(text):
            for delimiter in ['。\n', '。', '\n\n', '\n', '；', '；']:
                split_pos = text.rfind(delimiter, start, end)
                if split_pos > start:
                    end = split_pos + len(delimiter)
                    break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = end - overlap
    
    logger.info(f"✅ 文本分块完成: {len(text)} 字符 → {len(chunks)} 个块")
    return chunks

